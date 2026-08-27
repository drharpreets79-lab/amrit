from __future__ import annotations

import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "national_amr_ecosystem_plan.md"
OUTPUT = ROOT / "National_AMR_Ecosystem_Implementation_Blueprint_2026.docx"
FIGURE = ROOT / "national_amr_target_architecture.png"

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "087E8B"
GREEN = "3A7D44"
LIGHT_BLUE = "E8EEF5"
LIGHT_TEAL = "E7F3F4"
LIGHT_GRAY = "F4F6F9"
MID_GRAY = "D9E1E8"
TEXT = "1F2933"
MUTED = "5B6573"
WHITE = "FFFFFF"
GOLD = "B27A0B"
RED = "9B1C1C"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_table_borders(table, color=MID_GRAY, size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=120) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def paragraph_border_bottom(paragraph, color=BLUE, size="14", space="6") -> None:
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color=MUTED)


def make_numbering(doc: Document) -> tuple[int, int, int]:
    numbering = doc.part.numbering_part.element

    def next_id(tag: str, attr: str) -> int:
        vals = [int(n.get(qn(attr))) for n in numbering.findall(qn(tag)) if n.get(qn(attr))]
        return max(vals or [0]) + 1

    bullet_abstract = next_id("w:abstractNum", "w:abstractNumId")
    decimal_abstract = bullet_abstract + 1
    first_num = next_id("w:num", "w:numId")

    def add_abstract(abstract_id: int, fmt: str, text: str, font: str | None = None):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "280")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "290")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, ind, spacing])
        lvl.extend([start, num_fmt, lvl_text, suff, p_pr])
        if font:
            r_pr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), font)
            fonts.set(qn("w:hAnsi"), font)
            r_pr.append(fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)

    add_abstract(bullet_abstract, "bullet", "•", "Arial")
    add_abstract(decimal_abstract, "decimal", "%1.")

    def add_num(num_id: int, abstract_id: int):
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        ref = OxmlElement("w:abstractNumId")
        ref.set(qn("w:val"), str(abstract_id))
        num.append(ref)
        override = OxmlElement("w:lvlOverride")
        override.set(qn("w:ilvl"), "0")
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), "1")
        override.append(start_override)
        num.append(override)
        numbering.append(num)

    add_num(first_num, bullet_abstract)
    add_num(first_num + 1, decimal_abstract)
    return first_num, first_num + 1, decimal_abstract


def new_numbering_instance(doc: Document, abstract_id: int) -> int:
    numbering = doc.part.numbering_part.element
    vals = [
        int(n.get(qn("w:numId")))
        for n in numbering.findall(qn("w:num"))
        if n.get(qn("w:numId"))
    ]
    num_id = max(vals or [0]) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(abstract_id))
    num.append(ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    override.append(start_override)
    num.append(override)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    n = OxmlElement("w:numId")
    n.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, n])
    p_pr.append(num_pr)


def setup_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(TEXT)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.333

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Table Text" not in styles:
        table_style = styles.add_style("Table Text", WD_STYLE_TYPE.PARAGRAPH)
    else:
        table_style = styles["Table Text"]
    table_style.font.name = "Calibri"
    table_style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    table_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    table_style.font.size = Pt(9)
    table_style.font.color.rgb = rgb(TEXT)
    table_style.paragraph_format.space_before = Pt(0)
    table_style.paragraph_format.space_after = Pt(2)
    table_style.paragraph_format.line_spacing = 1.08

    if "Callout" not in styles:
        callout = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = styles["Callout"]
    callout.font.name = "Calibri"
    callout._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    callout._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    callout.font.size = Pt(10.5)
    callout.font.color.rgb = rgb(NAVY)
    callout.paragraph_format.space_before = Pt(8)
    callout.paragraph_format.space_after = Pt(10)
    callout.paragraph_format.left_indent = Inches(0.18)
    callout.paragraph_format.right_indent = Inches(0.12)
    callout.paragraph_format.line_spacing = 1.2


def add_hyperlink(paragraph, text: str, url: str, color=BLUE):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.extend([c, u])
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`|https?://\S+)")


def add_inline(paragraph, text: str, *, default_bold=False, default_color=TEXT, size=None) -> None:
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size=size, color=default_color, bold=default_bold)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=default_color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=(size or 10), color=DARK_BLUE)
        else:
            clean = token.rstrip(".,);]")
            suffix = token[len(clean):]
            add_hyperlink(paragraph, clean, clean)
            if suffix:
                run = paragraph.add_run(suffix)
                set_run_font(run, size=size, color=default_color)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size, color=default_color, bold=default_bold)


def create_architecture_figure(path: Path) -> None:
    width, height = 1800, 1100
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    font_paths = [
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial.ttf",
    ]
    fp = next((p for p in font_paths if Path(p).exists()), None)
    title_font = ImageFont.truetype(fp, 42) if fp else ImageFont.load_default()
    head_font = ImageFont.truetype(fp, 27) if fp else ImageFont.load_default()
    body_font = ImageFont.truetype(fp, 22) if fp else ImageFont.load_default()
    small_font = ImageFont.truetype(fp, 18) if fp else ImageFont.load_default()

    draw.text((60, 34), "National AMR Ecosystem — Federated Target Architecture", fill="#17365D", font=title_font)
    draw.text((60, 88), "Local systems retain operational records; governed data products and alerts drive national action.", fill="#5B6573", font=body_font)

    layers = [
        ("1  SECTOR EDGE PLANE", "#E7F3F4", "#087E8B", 160, 345),
        ("2  SECURE EXCHANGE PLANE", "#E8EEF5", "#2E74B5", 385, 565),
        ("3  NATIONAL COORDINATION PLANE", "#F4F6F9", "#17365D", 605, 820),
        ("4  KNOWLEDGE, RESEARCH & REPORTING", "#FFF4DA", "#B27A0B", 860, 1035),
    ]
    for label, fill, outline, y1, y2 in layers:
        draw.rounded_rectangle((45, y1, 1755, y2), radius=24, fill=fill, outline=outline, width=4)
        draw.text((70, y1 + 18), label, fill=outline, font=head_font)

    edge_boxes = [
        ("Human Lab", "AST • WHONET • alerts"),
        ("Stewardship / IPC", "AMC/AMU • AWaRe • HAI"),
        ("Animal / Fisheries", "AMR • AMU • biosecurity"),
        ("Food", "AMR • residues • tracing"),
        ("Environment", "water • effluent • genes"),
        ("Reference / Genomics", "confirm • sequence • cluster"),
    ]
    x = 70
    for title, subtitle in edge_boxes:
        box_w = 255
        draw.rounded_rectangle((x, 225, x + box_w, 320), radius=14, fill="white", outline="#087E8B", width=3)
        draw.text((x + 14, 239), title, fill="#17365D", font=body_font)
        draw.text((x + 14, 280), subtitle, fill="#5B6573", font=small_font)
        x += 278

    exchange_boxes = [
        ("Identity & PKI", 70, 455, 360),
        ("API gateway / event broker", 390, 455, 745),
        ("Schema & terminology registry", 775, 455, 1160),
        ("Policy • privacy • audit", 1190, 455, 1715),
    ]
    for text, x1, y1, x2 in exchange_boxes:
        draw.rounded_rectangle((x1, y1, x2, y1 + 75), radius=12, fill="white", outline="#2E74B5", width=3)
        tw = draw.textbbox((0, 0), text, font=body_font)[2]
        draw.text((x1 + (x2 - x1 - tw) / 2, y1 + 22), text, fill="#17365D", font=body_font)

    national_boxes = [
        ("Federated query & data products", 70, 690, 520),
        ("Quality • lineage • methods", 550, 690, 940),
        ("Alert • case • action • SLA", 970, 690, 1325),
        ("Role dashboards & NAP/SAP", 1355, 690, 1715),
    ]
    for text, x1, y1, x2 in national_boxes:
        draw.rounded_rectangle((x1, y1, x2, y1 + 82), radius=12, fill="white", outline="#17365D", width=3)
        tw = draw.textbbox((0, 0), text, font=body_font)[2]
        draw.text((x1 + (x2 - x1 - tw) / 2, y1 + 25), text, fill="#17365D", font=body_font)

    knowledge = ["Guidelines & playbooks", "Secure research enclave", "GLASS • ANIMUSE • InFARM", "Public transparency"]
    x = 90
    for text in knowledge:
        draw.rounded_rectangle((x, 930, x + 385, 1000), radius=12, fill="white", outline="#B27A0B", width=3)
        tw = draw.textbbox((0, 0), text, font=body_font)[2]
        draw.text((x + (385 - tw) / 2, 950), text, fill="#7A5A00", font=body_font)
        x += 420

    def arrow(y1, y2, color):
        x = 900
        draw.line((x, y1, x, y2), fill=color, width=7)
        draw.polygon([(x - 16, y2 - 22), (x + 16, y2 - 22), (x, y2)], fill=color)

    arrow(330, 385, "#087E8B")
    arrow(565, 605, "#2E74B5")
    arrow(820, 860, "#17365D")
    img.save(path, quality=95)


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    header = section.first_page_header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("OFFICE OF THE PRINCIPAL SCIENTIFIC ADVISER • NATIONAL AMR ECOSYSTEM")
    set_run_font(r, size=8.5, color=MUTED, bold=True)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(92)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(14)
    r = kicker.add_run("IMPLEMENTATION BLUEPRINT")
    set_run_font(r, size=11, color=TEAL, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    r = title.add_run("National AMR Ecosystem\nfor India")
    set_run_font(r, size=30, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    r = subtitle.add_run("A federated One Health programme built on AMRIT edge applications\nand a national web coordination platform")
    set_run_font(r, size=15, color=DARK_BLUE)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(24)
    paragraph_border_bottom(rule, color=TEAL, size="18", space="8")

    meta = doc.add_table(rows=5, cols=2)
    meta_data = [
        ("DOCUMENT PROFILE", "DETAILS"),
        ("Prepared for", "Office of the Principal Scientific Adviser and national/State/UT One Health stakeholders"),
        ("Planning horizon", "48 months from approval • first operational release in 9 months"),
        ("Version", "1.0 • 13 July 2026"),
        ("Status", "For consultation, costing, pilot approval and phased execution"),
    ]
    for ri, (row, (label, value)) in enumerate(zip(meta.rows, meta_data)):
        row.cells[0].text = label
        row.cells[1].text = value
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].style = doc.styles["Table Text"]
            for run in cell.paragraphs[0].runs:
                set_run_font(
                    run,
                    size=(8.8 if ri == 0 else 9.5),
                    color=(WHITE if ri == 0 else (NAVY if idx == 0 else TEXT)),
                    bold=(ri == 0 or idx == 0),
                )
    set_table_geometry(meta, [1800, 7560], indent_dxa=120)
    set_table_borders(meta, color="CCD6DF", size="6")
    set_repeat_table_header(meta.rows[0])
    for ri, row in enumerate(meta.rows):
        if ri == 0:
            for cell in row.cells:
                set_cell_shading(cell, NAVY)
        else:
            set_cell_shading(row.cells[0], LIGHT_BLUE)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(32)
    r = note.add_run("Decision-ready • standards-led • privacy-by-design • executable by work package")
    set_run_font(r, size=9.5, color=MUTED, italic=True)

    footer = section.first_page_footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("National AMR Ecosystem Implementation Blueprint")
    set_run_font(r, size=8.5, color=MUTED)
    doc.add_page_break()


def add_running_furniture(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("NATIONAL AMR ECOSYSTEM • IMPLEMENTATION BLUEPRINT")
    set_run_font(r, size=8.2, color=MUTED, bold=True)
    paragraph_border_bottom(p, color=MID_GRAY, size="6", space="3")
    footer = section.footer
    p = footer.paragraphs[0]
    p.clear()
    left = p.add_run("Government consultation draft • Version 1.0")
    set_run_font(left, size=8.5, color=MUTED)
    tab = OxmlElement("w:tab")
    left._r.addnext(tab)
    add_page_number(p)


def add_contents(doc: Document, headings: list[str]) -> None:
    p = doc.add_paragraph("Contents", style="Heading 1")
    p.paragraph_format.page_break_before = False
    intro = doc.add_paragraph("The document is organised as a decision and delivery sequence. Word’s Navigation Pane can be used to jump between numbered sections.")
    intro.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for text in headings:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        set_run_font(r, size=10.3, color=DARK_BLUE)
    doc.add_page_break()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r"-+", c.replace(":", "")) for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def table_widths(rows: list[list[str]]) -> list[int]:
    cols = len(rows[0])
    first_header = rows[0][0].strip().lower() if rows and rows[0] else ""
    if cols == 2:
        return [2450, 6910]
    if cols == 3:
        return [1800, 3500, 4060]
    if cols == 4:
        if first_header == "risk":
            return [1850, 1700, 2910, 2900]
        if first_header == "domain":
            return [1450, 2250, 3500, 2160]
        return [1200, 2450, 3000, 2710]
    if cols == 5:
        return [850, 2150, 1450, 2550, 2360]
    weights = []
    for c in range(cols):
        max_len = max(len(r[c]) if c < len(r) else 0 for r in rows)
        weights.append(max(8, min(max_len, 45)))
    total = sum(weights)
    widths = [round(9360 * w / total) for w in weights]
    widths[-1] += 9360 - sum(widths)
    return widths


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    set_table_geometry(table, table_widths(rows), indent_dxa=120)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for ri, row in enumerate(rows):
        prevent_row_split(table.rows[ri])
        for ci in range(cols):
            cell = table.rows[ri].cells[ci]
            text = row[ci] if ci < len(row) else ""
            p = cell.paragraphs[0]
            p.style = doc.styles["Table Text"]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, text, default_bold=(ri == 0), default_color=(WHITE if ri == 0 else TEXT), size=(8.7 if len(rows) > 20 else 9))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ri == 0:
                set_cell_shading(cell, NAVY)
            elif ri % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def build() -> None:
    create_architecture_figure(FIGURE)
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    body_start = next(i for i, line in enumerate(lines) if line.startswith("> **EXECUTIVE DECISION."))
    headings = [line[2:].strip() for line in lines if line.startswith("# ") and not line.startswith("##")]
    headings += [line[3:].strip() for line in lines if line.startswith("## ") and re.match(r"(?:\d+\.|Appendix)", line[3:].strip())]

    doc = Document()
    setup_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_cover(doc)
    add_running_furniture(section)

    contents = [
        "1. Purpose and intended use", "2. Executive assessment", "3. National mission outcomes",
        "4. Target operating model", "5. Target technical architecture", "6. Product portfolio",
        "7. Surveillance and analytics", "8. Governance, privacy and cybersecurity",
        "9. Implementation roadmap", "10. Executable backlog", "11. Testing and assurance",
        "12. Adoption and support", "13. Procurement and delivery", "14. Monitoring and evaluation",
        "15. Risk register", "16. First 100 days", "17. Leadership decision gates",
        "18. Definition of success", "Appendices A–E",
    ]
    add_contents(doc, contents)

    bullet_id, decimal_id, decimal_abstract = make_numbering(doc)
    i = body_start
    architecture_inserted = False
    last_kind = None
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            last_kind = None
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            last_kind = None
            continue
        if line.startswith("# "):
            last_kind = None
            i += 1
            continue
        if line.startswith("## "):
            text = line[3:].strip()
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, text, default_bold=True, default_color=BLUE, size=16)
            last_kind = None
            i += 1
            continue
        if line.startswith("### "):
            text = line[4:].strip()
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, text, default_bold=True, default_color=BLUE, size=13)
            if text == "5.1 Logical architecture" and not architecture_inserted:
                intro = doc.add_paragraph("The figure below shows how the existing desktop/server combination becomes a federated national ecosystem while retaining sector ownership of operational data.")
                intro.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                pic = doc.add_picture(str(FIGURE), width=Inches(6.45))
                pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap = doc.add_paragraph("Figure 1. Federated target architecture for the National AMR Ecosystem")
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.paragraph_format.space_after = Pt(10)
                for run in cap.runs:
                    set_run_font(run, size=9, color=MUTED, italic=True)
                pic._inline.docPr.set("title", "Federated target architecture")
                pic._inline.docPr.set(
                    "descr",
                    "Four-layer diagram showing sector edge applications, secure exchange services, national coordination services, and knowledge, research and reporting outputs.",
                )
                architecture_inserted = True
            last_kind = None
            i += 1
            continue
        if line.startswith("**") and line.endswith("**") and len(line) < 120:
            p = doc.add_paragraph(style="Heading 3")
            add_inline(p, line, default_bold=True, default_color=DARK_BLUE, size=12)
            last_kind = None
            i += 1
            continue
        if line.startswith("> "):
            p = doc.add_paragraph(style="Callout")
            add_inline(p, line[2:], default_color=NAVY, size=10.5)
            p_pr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), LIGHT_TEAL)
            p_pr.append(shd)
            left_border = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "24")
            left.set(qn("w:space"), "8")
            left.set(qn("w:color"), TEAL)
            left_border.append(left)
            p_pr.append(left_border)
            last_kind = None
            i += 1
            continue
        if re.match(r"^- ", line):
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            apply_num(p, bullet_id)
            add_inline(p, line[2:].strip())
            last_kind = "bullet"
            i += 1
            continue
        if re.match(r"^\d+\. ", line):
            if last_kind != "decimal":
                decimal_id = new_numbering_instance(doc, decimal_abstract)
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            apply_num(p, decimal_id)
            add_inline(p, re.sub(r"^\d+\. ", "", line))
            last_kind = "decimal"
            i += 1
            continue

        # Join adjacent prose lines until a structural marker.
        parts = [line]
        j = i + 1
        while j < len(lines):
            nxt = lines[j].strip()
            if not nxt or nxt.startswith(("#", "|", ">", "- ")) or re.match(r"^\d+\. ", nxt):
                break
            if nxt.startswith("**") and nxt.endswith("**") and len(nxt) < 120:
                break
            parts.append(nxt)
            j += 1
        text = " ".join(parts)
        p = doc.add_paragraph()
        if text.startswith("`") and text.endswith("`"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(p, text)
        last_kind = None
        i = j

    # Footer tab stop for all sections.
    for sec in doc.sections:
        add_running_furniture(sec)
        sec.page_width = Inches(8.5)
        sec.page_height = Inches(11)
        sec.top_margin = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)
        sec.header_distance = Inches(0.492)
        sec.footer_distance = Inches(0.492)

    # Core properties and update fields on open.
    doc.core_properties.title = "National AMR Ecosystem for India — Implementation Blueprint"
    doc.core_properties.subject = "Federated One Health AMR surveillance, stewardship and response programme"
    doc.core_properties.author = "Prepared with Codex for consultation"
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
